from docx import Document as DocxDocument
from spire.doc import Document as SpireDocument
from spire.doc.common import *
import re
import pandas as pd
import requests
from datetime import datetime, timedelta
import pytz
from tabulate import tabulate
from textwrap import fill
from nltk import word_tokenize, pos_tag

def read_tables_from_docx(file_path):
    doc = DocxDocument(file_path)
    all_table_data = []

    # Read all tables in the document
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        all_table_data.append(table_data)

    return all_table_data

def extract_sorted_unique_versions(file_path):
    doc = SpireDocument()
    doc.LoadFromFile(file_path)
    text = doc.GetText()
    pattern = r"\b\d+\.\d+\.\d+\.\d+\b"
    matches = re.findall(pattern, text)
    unique_matches = sorted(set(matches), key=lambda s: list(map(int, s.split('.'))))
    return unique_matches

def process_all_tables(all_table_data, value_list):
    total_rows = sum(len(table) - 1 for table in all_table_data)  # Exclude header rows from each table
    num_values = len(value_list)

    if total_rows != num_values:
        return "The number of values does not match the total number of rows across all tables"

    current_value_index = 0

    for table_data in all_table_data:
        header = table_data[0]
        if "Step #" in header:
            step_index = header.index("Step #")
        else:
            return "No 'Step #' column found in one of the tables"

        for i in range(1, len(table_data)):
            table_data[i][step_index] = value_list[current_value_index]
            current_value_index += 1

    return all_table_data

def extract_tables_from_docx(docx_file):
    doc = DocxDocument(docx_file)
    tables_data = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        df = pd.DataFrame(table_data[1:], columns=table_data[0])
        tables_data.append(df)
    return tables_data


def check_screenshot_and_attachment(df):
    results = []
    hyperlink_pattern = re.compile(r'(http[s]?://\S+|www\.\S+|\[[^\]]+\]|\S+/[\w./-]+)')

    for step, test_procedure, actual_result in zip(df["Step #"], df["Test Procedure"], df["Actual Result"]):
        if re.search(r"capture screenshot\(s\)", test_procedure, re.IGNORECASE):
            hyperlink_present = hyperlink_pattern.findall(actual_result)
            if hyperlink_present:
                result_message = f"{step}-- has hyperlink(s): {', '.join(hyperlink_present)}."
                results.append(result_message)
            else:
                result_message = f"{step}-- has no hyperlink."
                results.append(result_message)
    
    return results


def extract_timestamp_and_step(df):
    step_column = df["Step #"]
    executed_column = df["Executed By & Date"]
    timestamp_pattern = r"\d{2}-[a-zA-Z]{3}-\d{4} \d{2}:\d{2}:\d{2} [AP]M \([A-Z]+\)"
    timestamps_and_steps = []
    prev_timestamp = None
    ist = pytz.timezone('Asia/Kolkata')
    for step, executed_str in zip(step_column, executed_column):
        timestamp_match = re.search(timestamp_pattern, executed_str)
        if timestamp_match:
            current_timestamp = timestamp_match.group(0)
            if prev_timestamp:
                try:
                    current_time = datetime.strptime(re.sub(r" \([A-Z]+\)", "", current_timestamp), "%d-%b-%Y %I:%M:%S %p")
                    current_time_ist = ist.localize(current_time)
                    
                    prev_time = datetime.strptime(re.sub(r" \([A-Z]+\)", "", prev_timestamp), "%d-%b-%Y %I:%M:%S %p")
                    prev_time_ist = ist.localize(prev_time)
                    
                    time_diff = current_time_ist - prev_time_ist
                    
                    if time_diff > timedelta(hours=1):
                        timestamps_and_steps.append(f"{step} - Time difference with previous step exceeds 1 hour")
                except ValueError as e:
                    timestamps_and_steps.append(f"Timestamp format error: {e}")
                   
            prev_timestamp = current_timestamp
    
    return timestamps_and_steps

def wrap_text(text, width):
    """Wrap text to a specified width."""
    return fill(text, width=width)

def check_step_results(df):
    step_column = df["Step #"]
    expected_column = df["Expected Result"]
    actual_column = df["Actual Result"]
    pass_fail_column = df["Pass/Fail"]

    results = []

    for step, expected, actual, result in zip(step_column, expected_column, actual_column, pass_fail_column):
        if result.lower() == "fail":
            wrapped_expected = wrap_text(expected, 20)
            wrapped_actual = wrap_text(actual, 20)
            results.append((step, "Fail", wrapped_expected, wrapped_actual))

    # Define table headers
    headers = ["Step #", "Status", "Expected Result", "Actual Result"]

    if results:
        table = tabulate(results, headers=headers, tablefmt="grid")
        print(table)
    else:
        print("No Fail on this Table")

def is_present_tense(sentence):
    tokens = word_tokenize(sentence)
    tags = pos_tag(tokens)
    for word, tag in tags:
        if tag in {"VB", "VBP", "VBZ"}:
            return True
    return False 

def check_actuals_present_tense(df):
    results = []
    for step, actual in zip(df["Step #"], df["Actual Result"]):
        tense_status = "Present Tense"
        if actual:  # Check if there is a sentence
            tense_status = "Present Tense" if is_present_tense(actual) else "Not Present Tense"
            results.append(f"{step}-{tense_status}")
        else:
            results.append(f"{step}-No sentence")
    return results

file_path = "Step-1.docx"

# Extract sorted unique versions for updating "Step #" column
sorted_unique_versions = extract_sorted_unique_versions(file_path)

# Read all tables from DOCX and process them with the new "Step #" values
all_table_data = read_tables_from_docx(file_path)
processed_table_data = process_all_tables(all_table_data, sorted_unique_versions)

# Extract other details from tables and store them in DataFrame
tables = extract_tables_from_docx(file_path)

# Print DataFrame with updated "Step #" and other details
for i, table in enumerate(tables):
    # Update "Step #" column in the DataFrame
    table["Step #"] = [row[0] for row in processed_table_data[i][1:]]
    
    print(f"Table {i+1}:")

     # Check screenshot and attachment details
    hyperlink_details = check_screenshot_and_attachment(table)
    print("----------------Hyperlink Details:-------------")
    for detail in hyperlink_details:
        print(detail)
    
    
